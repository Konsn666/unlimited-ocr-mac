"""Export OCR results to Word (.docx) with formatting."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_docx(pages: list[dict], output_path: Path, export_mode: str = "original") -> Path:
    """Build a .docx from structured page data.

    export_mode: "original" | "translated" | "bilingual"
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(12)

    for pi, page in enumerate(pages):
        page_num = page["page_num"]
        blocks = page.get("blocks", [])
        detections = page.get("detections", [])
        translations = page.get("translations", [])

        # Build translation map: detection_index -> translated text
        trans_map = {}
        if translations:
            for t in translations:
                if "index" in t and "translated" in t:
                    trans_map[t["index"]] = t["translated"]

        if pi > 0:
            doc.add_page_break()

        for block in blocks:
            bt = block["block_type"]
            text = block.get("text", "")
            if not text.strip():
                continue

            # Get detection index (use first detection index in block)
            det_idx = _get_block_det_index(block, detections)

            # Determine which text(s) to output
            original = text
            translated = trans_map.get(det_idx, "") if trans_map else ""

            if bt == "heading":
                level = block.get("level", 2)
                if export_mode == "translated" and translated:
                    _add_heading(doc, translated, level)
                elif export_mode == "bilingual" and translated:
                    _add_heading(doc, original, level)
                    _add_translation_para(doc, translated, is_heading=True)
                else:
                    _add_heading(doc, original, level)

            elif bt == "paragraph":
                lines = [l for l in text.split("\n") if l.strip()]
                for line in lines:
                    if export_mode == "translated" and translated:
                        doc.add_paragraph(translated)
                    elif export_mode == "bilingual" and translated:
                        doc.add_paragraph(original)
                        _add_translation_para(doc, translated)
                    else:
                        doc.add_paragraph(original)

            elif bt == "table":
                table_text = translated if (export_mode == "translated" and translated) else original
                if export_mode == "bilingual" and translated:
                    p = doc.add_paragraph()
                    run = p.add_run(original)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    _add_translation_para(doc, translated)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(table_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)

            elif bt == "image":
                p = doc.add_paragraph()
                run = p.add_run("[图片]")
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.italic = True

            elif bt == "page_number":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
    return output_path


def _add_heading(doc, text, level):
    h = doc.add_heading(text, level=min(level, 3))
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_translation_para(doc, text, is_heading=False):
    """Add a translated text paragraph with distinct styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(79, 70, 229)  # Indigo
    if is_heading:
        run.font.size = Pt(13)
        run.bold = True
    else:
        run.font.size = Pt(11)
    # Left border via indentation
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _get_block_det_index(block, detections):
    """Find the first detection index that corresponds to this block."""
    block_text = block.get("text", "")
    block_type = block.get("block_type", "")
    for i, det in enumerate(detections):
        if det.get("text") == block_text and det.get("type") in ("text", "title", "table", block_type):
            return i
    return -1
